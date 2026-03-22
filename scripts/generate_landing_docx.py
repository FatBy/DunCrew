# -*- coding: utf-8 -*-
"""
DunCrew 产品页文案 + 截图 Word 文档生成脚本
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), '..', 'screenshots')
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '..', 'DunCrew产品页素材.docx')


def set_run_font(run, name='微软雅黑', size=None, bold=False, color=None):
    run.font.name = name
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体
    from docx.oxml.ns import qn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_heading_styled(doc, text, level=1, color=(30, 30, 30)):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_run_font(run, size={1: 22, 2: 18, 3: 14}.get(level, 12), bold=True, color=color)
    return h


def add_para(doc, text, size=11, color=(60, 60, 60), bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_image_centered(doc, img_path, width_inches=6.0, caption=''):
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=9, color=(150, 150, 150))
        cap.paragraph_format.space_after = Pt(4)


def build_docx():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    from docx.oxml.ns import qn
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()

    add_para(doc, 'DunCrew', size=36, bold=True, color=(30, 30, 30), align='center', space_after=8)
    add_para(doc, '给你的 AI 一个工位', size=18, color=(80, 80, 80), align='center', space_after=16)
    add_para(doc, 'duncrew.cn 产品落地页素材包', size=12, color=(130, 130, 130), align='center', space_after=4)
    add_para(doc, '文案 + 产品截图', size=12, color=(130, 130, 130), align='center')

    # ========== 产品定位 ==========
    doc.add_page_break()
    add_heading_styled(doc, '产品定位', level=1)
    add_para(doc, '一句话：给你的 AI 一个工位', size=16, bold=True, color=(60, 60, 60), space_after=12)
    add_para(doc, (
        'DunCrew 不是又一个聊天机器人。\n'
        '它是一套运行在你电脑上的 AI 操作系统 -- 你可以给 AI 分配"工位"，'
        '让它们各司其职，越用越专业。\n\n'
        '不上传数据，不依赖云端，所有能力都长在你自己的电脑里。'
    ), size=11, color=(80, 80, 80))

    # ========== 6 张功能卡 ==========
    cards = [
        {
            'num': '01',
            'title': 'Nexus 专家体系',
            'subtitle': '养虾，不如孵蛋',
            'body': (
                '大多数 AI 工具是"现成的虾" -- 买来就用，但永远不会变。\n'
                'DunCrew 里的每个 Nexus 是一颗"蛋"：\n'
                '  - 刚孵出来什么都不太会\n'
                '  - 用得越多，越懂你的需求\n'
                '  - 每个 Nexus 会发展出自己的专长\n'
                '  - 最终变成只属于你的专家团队\n\n'
                '你不是在"用工具"，你是在"养团队"。'
            ),
            'screenshot': '01-world-dashboard.png',
            'caption': 'World 画布 -- 你的 AI 专家团队一览',
        },
        {
            'num': '02',
            'title': '自我反思',
            'subtitle': '摔过的坑，自己填上',
            'body': (
                '普通 AI 犯了错，下次还犯。\n'
                'DunCrew 的 Nexus 不一样：\n'
                '  - 每次任务结束，自动复盘\n'
                '  - 哪里做得好、哪里翻车，它自己写总结\n'
                '  - 下次遇到类似情况，会自动绕开老坑\n\n'
                '不需要你反复纠正，它自己就在进步。'
            ),
            'screenshot': '03-chat-interface.png',
            'caption': 'AI 对话界面 -- 结构化回复 + 选项引导',
        },
        {
            'num': '03',
            'title': '真正的记忆',
            'subtitle': '它真的记得你',
            'body': (
                '大多数 AI 聊完就忘，DunCrew 有三层记忆：\n'
                '  - 短期记忆：这次对话里发生了什么\n'
                '  - 长期记忆：上个月帮你做过什么\n'
                '  - 共享记忆：一个 Nexus 学到的经验，其他 Nexus 也能用\n\n'
                '"上次那个报告的格式，再来一份" -- 它真的能做到。'
            ),
            'screenshot': None,
            'caption': '',
        },
        {
            'num': '04',
            'title': '本地运行',
            'subtitle': '门都不出',
            'body': (
                '你的文件、对话、记忆 -- 全都在你自己电脑上。\n'
                '  - 不上传到任何云端服务器\n'
                '  - 不需要注册账号\n'
                '  - 断网了？本地模型照样跑\n\n'
                '隐私不是功能，是底线。'
            ),
            'screenshot': '04-soul-tower.png',
            'caption': 'Soul 灵魂塔 -- MBTI 人格系统，独一无二的 AI 身份',
        },
        {
            'num': '05',
            'title': '兼容 OpenClaw',
            'subtitle': '熟悉的味道，不一样的灵魂',
            'body': (
                '如果你用过 OpenAI 的 GPTs，上手零门槛：\n'
                '  - 一键导入 OpenClaw 配置\n'
                '  - 但这里的 AI 会成长、会反思、有记忆\n'
                '  - 同样的配置，在 DunCrew 里活过来\n\n'
                '不是替代品，是进化版。'
            ),
            'screenshot': '02-nexus-detail.png',
            'caption': 'Nexus 详情面板 -- 成长评分、成就系统、能力维度',
        },
        {
            'num': '06',
            'title': '三步开始',
            'subtitle': '下载，打开，开聊',
            'body': (
                '第一步：下载 DunCrew 安装包 (Windows)\n'
                '第二步：打开后，在设置里填入你的 API Key\n'
                '         (支持 OpenAI / Claude / 国产大模型)\n'
                '第三步：选一个 Nexus 开始对话，或者自己创建一个\n\n'
                '不需要懂代码，不需要配环境。\n'
                '下载、打开、开聊。就这么简单。'
            ),
            'screenshot': '01-world-dashboard.png',
            'caption': 'World 全景 -- 创建你的第一个 Nexus 专家',
        },
    ]

    for card in cards:
        doc.add_page_break()

        # 编号 + 标题
        p = doc.add_paragraph()
        r_num = p.add_run(f'{card["num"]}  ')
        set_run_font(r_num, size=28, bold=True, color=(200, 200, 200))
        r_title = p.add_run(card['title'])
        set_run_font(r_title, size=20, bold=True, color=(30, 30, 30))

        # 副标题
        add_para(doc, card['subtitle'], size=14, bold=True, color=(100, 100, 100), space_after=8)

        # 正文
        add_para(doc, card['body'], size=11, color=(60, 60, 60), space_after=12)

        # 截图
        if card['screenshot']:
            img_path = os.path.join(SCREENSHOTS_DIR, card['screenshot'])
            add_image_centered(doc, img_path, width_inches=6.2, caption=card['caption'])

    # ========== 商业描述页 ==========
    doc.add_page_break()
    add_heading_styled(doc, '阿里云建站 -- 商业描述', level=1)

    biz_items = [
        ('品牌名称', 'DunCrew'),
        ('域名', 'duncrew.cn'),
        ('定位', '给你的 AI 一个工位'),
        ('产品类型', '本地运行的 AI 操作系统（桌面应用）'),
        ('目标用户', '需要频繁使用 AI 完成工作的个人和小团队\n（内容创作者、运营、研究员、自由职业者等）'),
    ]

    for label, value in biz_items:
        p = doc.add_paragraph()
        r_label = p.add_run(f'{label}：')
        set_run_font(r_label, size=11, bold=True, color=(30, 30, 30))
        r_value = p.add_run(value)
        set_run_font(r_value, size=11, color=(60, 60, 60))
        p.paragraph_format.space_after = Pt(4)

    add_para(doc, '', space_after=8)
    add_para(doc, '核心差异化：', size=12, bold=True, color=(30, 30, 30), space_after=4)

    diffs = [
        'AI 会成长 -- 越用越专业（Nexus 专家体系 + 等级系统）',
        'AI 会反思 -- 犯了错自己复盘（Reflexion 机制）',
        'AI 有记忆 -- 记得你的偏好和历史（三层记忆架构）',
        '完全本地 -- 数据不出你的电脑（隐私优先）',
        '有个性 -- MBTI 人格系统，不是冷冰冰的工具',
    ]
    for i, d in enumerate(diffs, 1):
        add_para(doc, f'{i}. {d}', size=11, color=(60, 60, 60), space_after=3)

    add_para(doc, '', space_after=8)
    add_para(doc, '分发方式：Windows 桌面安装包（一键安装）', size=11, color=(60, 60, 60), space_after=3)
    add_para(doc, '商业模式：开源免费，用户自带 API Key', size=11, color=(60, 60, 60))

    # 保存
    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f'Word 文档已生成: {OUTPUT_PATH}')
    print(f'文件大小: {size_kb:.0f} KB ({size_kb/1024:.1f} MB)')


if __name__ == '__main__':
    build_docx()
